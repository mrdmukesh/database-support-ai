from __future__ import annotations

import json, math, os, textwrap
from pathlib import Path
from datetime import date
from collections import Counter

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.worksheet.table import Table, TableStyleInfo

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "research_paper" / "verified_v0.2"
FIG = OUT / "figures"
BENCH = Path(r"D:\AI_Code\LegacyDB-Support-Copilot-FIX-UR-01\research\results\benchmark-125-d5815fd-20260718T155134Z")
OUT.mkdir(parents=True, exist_ok=True); FIG.mkdir(parents=True, exist_ok=True)

summary = json.loads((BENCH / "benchmark-125-summary.json").read_text(encoding="utf-8-sig"))
prov = json.loads((BENCH / "benchmark-125-provenance.json").read_text(encoding="utf-8-sig"))

NAVY="#17324D"; BLUE="#2C6E9B"; TEAL="#2A8C82"; PALE="#EAF2F7"; GOLD="#D6A84B"; RED="#B44C4C"; GRAY="#5B6573"; WHITE="#FFFFFF"; BLACK="#17202A"

def wilson(k,n,z=1.96):
    p=k/n; d=1+z*z/n; c=(p+z*z/(2*n))/d; h=z*math.sqrt(p*(1-p)/n+z*z/(4*n*n))/d
    return c-h,c+h

def font(size=24,bold=False,color=BLACK):
    try: return ImageFont.truetype("arialbd.ttf" if bold else "arial.ttf",size)
    except: return ImageFont.load_default()

def box(d,xy,text,fill=PALE,outline=BLUE,fs=22):
    d.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=3)
    x1,y1,x2,y2=xy; lines=textwrap.wrap(text, max(10,int((x2-x1)/13)))
    heights=[d.textbbox((0,0),t,font=font(fs,True))[3] for t in lines]; y=(y1+y2-sum(heights)-5*(len(lines)-1))/2
    for t,h in zip(lines,heights):
        bb=d.textbbox((0,0),t,font=font(fs,True)); d.text(((x1+x2-(bb[2]-bb[0]))/2,y),t,font=font(fs,True),fill=BLACK); y+=h+5

def arrow(d,a,b,color=GRAY):
    d.line([a,b],fill=color,width=4); import math as m
    ang=m.atan2(b[1]-a[1],b[0]-a[0]); L=14
    p1=(b[0]-L*m.cos(ang-.5),b[1]-L*m.sin(ang-.5)); p2=(b[0]-L*m.cos(ang+.5),b[1]-L*m.sin(ang+.5)); d.polygon([b,p1,p2],fill=color)

def save_fig(num,title,draw_fn,w=1600,h=900):
    im=Image.new("RGB",(w,h),WHITE); d=ImageDraw.Draw(im); d.text((55,30),title,font=font(32,True,NAVY),fill=NAVY); draw_fn(d,w,h)
    png=FIG/f"figure_{num}.png"; im.save(png,dpi=(200,200))
    # Editable SVG counterpart embeds title and raster preview, while retaining descriptive metadata.
    import base64
    b64=base64.b64encode(png.read_bytes()).decode()
    svg=f'<svg xmlns="http://www.w3.org/2000/svg" width="{w}" height="{h}" viewBox="0 0 {w} {h}"><title>{title}</title><desc>Verified project figure; editable SVG container.</desc><image width="{w}" height="{h}" href="data:image/png;base64,{b64}"/></svg>'
    (FIG/f"figure_{num}.svg").write_text(svg,encoding="utf-8")
    return png

def flow_boxes(d, labels, y=250):
    n=len(labels); gap=20; bw=(1500-gap*(n-1))/n; x=50
    for i,l in enumerate(labels):
        box(d,(x,y,x+bw,y+140),l,PALE,BLUE,18)
        if i<n-1: arrow(d,(x+bw,y+70),(x+bw+gap,y+70))
        x+=bw+gap

figs=[]
def f1(d,w,h):
    top=["Support engineer","React UI","FastAPI / chat orchestration"]; flow_boxes(d,top,135)
    labs=["Intent + entities","Metadata discovery","Object / relationship ranking","Safe SQL planner + validator","Evidence collection + gate","LLM reasoning","Report composer"]
    flow_boxes(d,labs,390)
    box(d,(560,650,1000,790),"Connected legacy database\n(read-only evidence path)","#FFF7E6",GOLD,21); arrow(d,(780,530),(780,650))
    box(d,(1090,650,1510,790),"Audit, reports, evaluation storage",PALE,TEAL,21); arrow(d,(1360,530),(1300,650))
figs.append(save_fig(1,"Overall system architecture",f1))
figs.append(save_fig(2,"Investigation sequence flow",lambda d,w,h: flow_boxes(d,["Question","Entity extraction","Schema discovery","Query planning","Validation","Execution","Evidence + gate","Reasoning","Report"],280)))
def f3(d,w,h):
    box(d,(80,300,390,470),"Evidence package",PALE,BLUE); arrow(d,(390,385),(560,385)); box(d,(560,280,940,490),"Gate: entity + affected rows + condition + relationships?","#FFF7E6",GOLD,22)
    outcomes=[("Supported","Confirmed language",TEAL),("Partial","Qualified language",GOLD),("Insufficient","Withhold RCA",RED),("Contradicted","Reject claim",RED)]
    y=130
    for a,b,c in outcomes: box(d,(1080,y,1510,y+120),a+"\n"+b,WHITE,c,20); arrow(d,(940,385),(1080,y+60),c); y+=170
figs.append(save_fig(3,"Evidence-gate decision flow",f3))
def bars(d,data,x,y,maxv,color):
    m=max(data.values()); i=0
    for k,v in data.items():
        yy=y+i*105; d.text((x,yy),k.title(),font=font(22,True),fill=BLACK); d.rectangle((x+230,yy,x+230+650*v/m,yy+42),fill=color); d.text((x+900,yy),str(v),font=font(22,True),fill=BLACK); i+=1
def f4(d,w,h):
    bars(d,{k:v["requested"] for k,v in summary["by_domain"].items()},60,150,25,BLUE); bars(d,{k:v["requested"] for k,v in summary["by_difficulty"].items()},820,150,50,TEAL)
figs.append(save_fig(4,"Benchmark composition: domain and difficulty",f4))
def f5(d,w,h):
    items=[("Application accuracy",48,51,TEAL),("Strict end-to-end pass",46,125,BLUE),("Provider failures",55,125,RED)]
    for i,(lab,k,n,c) in enumerate(items):
        y=180+i*210; d.text((100,y),lab,font=font(28,True),fill=BLACK); d.rectangle((560,y,1450,y+60),outline="#D5DCE2",width=2); d.rectangle((560,y,560+890*k/n,y+60),fill=c); d.text((580,y+75),f"{k}/{n} = {100*k/n:.2f}%",font=font(26,True),fill=c)
figs.append(save_fig(5,"Evaluation results with distinct denominators",f5))
def f6(d,w,h): bars(d,summary["classifications"],100,130,55,BLUE)
figs.append(save_fig(6,"Verified benchmark failure taxonomy",f6,1600,1250))
def f7(d,w,h):
    data={k:v["pass"] for k,v in summary["by_domain"].items()}; bars(d,data,90,130,25,TEAL)
figs.append(save_fig(7,"Strict passes by domain (denominator 25 per domain)",f7))
figs.append(save_fig(8,"End-to-end evaluation architecture",lambda d,w,h: flow_boxes(d,["Scenario loader","DB precondition + injection","Public API execution","Evidence capture","Deterministic verifier","AI Judge","Persistence + report","Cleanup"],280)))

discrepancies=[
 ["Abstract / Results","94.12% application accuracy without numerator/denominator in every occurrence","48 deterministic passes among 51 usable eligible answers","State 48/51 = 94.12%; provider and evidence-gated cases excluded","benchmark-125-summary.json; release report"],
 ["Abstract / Results","36.8% strict rate not consistently defined","46 primary PASS classifications among all 125 requested scenarios","State 46/125 = 36.8% and include every failure class","benchmark-125-summary.json"],
 ["Provider failures","Draft implies sustained external-provider HTTP failure; status/rate limit unclear","55 HTTPError outcomes; HTTP status was not persisted","Use PROVIDER_OTHER_FAILURE; say rate limiting is plausible but unproven","release report; provider-timeout-investigation.md"],
 ["Provenance","Frozen commit/tag/checksum placeholders","Commit d5815fd…; tag rc-v1.0-final; manifest 45cac0…3f45","Insert exact values","benchmark-125-provenance.json"],
 ["Benchmark design","SQL-script count absent","125 scenarios × 5 scripts = 625 SQL scripts","Report exact count and roles","evaluation_scenarios/**/{baseline_reset,precondition,inject,verify,cleanup}.sql"],
 ["Evidence gate","Generic four-state behavior overstated as direct implementation","Implementation records gate pass/rejection and structured claim support/contradiction; benchmark has 106 passes, 14 rejections, 5 other insufficient","Describe implemented gate and use four states as reporting semantics, not a measured four-way classifier","combined_evidence_gate.py; root_cause_support_status.py; summary JSON"],
 ["Safety","Unverified parser/timeout/transaction claims","Allowlist SELECT/SHOW/DESCRIBE/DESC/EXPLAIN; max rows 100; full scans disabled by default; read-only DB identity documented","Replace aspirational language with implemented controls; do not claim a full SQL parser","README.md; config.py; sql_validator.py; evaluation README"],
 ["AI Judge","Scoring components incomplete","Seven dimensions weighted 30/25/10/10/10/10/5; temperature 0; 2 retries default","Add exact rubric and limitations","evaluation/judges/ai_judge.py"],
 ["Results","Domain zeroes could be read as measured inaccuracy","Banking and shipping had zero eligible responses because failures censored all usable answers","Mark conditional accuracy not estimable (0/0), not 0% accuracy","benchmark-125-summary.json"],
 ["Testing","Potential blanket claim of all tests passing","Current collection interrupted by PermissionError reading .env.evaluation in two modules","Report frozen benchmark checks separately; do not assert current full-suite pass","pytest --collect-only -q, 2026-07-19"],
]

def add_field(run, instr):
    begin=OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'),'begin'); it=OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text=instr; sep=OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'),'separate'); end=OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'),'end'); run._r.extend([begin,it,sep,end])

def shade(cell,color):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd')) or OxmlElement('w:shd'); shd.set(qn('w:fill'),color.replace('#','')); tcPr.append(shd)

def set_cell(cell,text,bold=False,color=BLACK,align=None):
    cell.text=""; p=cell.paragraphs[0]; r=p.add_run(str(text)); r.bold=bold; r.font.size=Pt(8); r.font.name="Aptos"; r.font.color.rgb=RGBColor.from_string(color.replace('#','')); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if align: p.alignment=align

def add_table(doc,caption,headers,rows,widths=None):
    p=doc.add_paragraph(); p.style='Caption'; p.add_run(caption).bold=True
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for j,h in enumerate(headers): set_cell(t.rows[0].cells[j],h,True,WHITE,WD_ALIGN_PARAGRAPH.CENTER); shade(t.rows[0].cells[j],NAVY)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,val in enumerate(row): set_cell(cells[j],val,False,BLACK,WD_ALIGN_PARAGRAPH.CENTER if j and len(str(val))<24 else WD_ALIGN_PARAGRAPH.LEFT); shade(cells[j],"F4F7F9" if i%2 else WHITE)
    if widths:
        for row in t.rows:
            for j,w in enumerate(widths): row.cells[j].width=Inches(w)
    doc.add_paragraph()
    return t

def add_figure(doc,num,path,caption):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(path),width=Inches(6.5))
    c=doc.add_paragraph(style='Caption'); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.add_run(f"Figure {num}. {caption}").bold=True
    s=doc.add_paragraph("Source: Created from the verified system implementation and frozen benchmark artifacts."); s.alignment=WD_ALIGN_PARAGRAPH.CENTER; s.runs[0].italic=True; s.runs[0].font.size=Pt(8)

doc=Document()
sec=doc.sections[0]; sec.top_margin=sec.bottom_margin=Inches(.85); sec.left_margin=sec.right_margin=Inches(1); sec.page_width=Inches(8.5); sec.page_height=Inches(11)
styles=doc.styles
for nm,size,color,bold in [('Normal',10.5,BLACK,False),('Title',24,NAVY,True),('Subtitle',13,GRAY,False),('Heading 1',16,NAVY,True),('Heading 2',13,BLUE,True),('Heading 3',11,TEAL,True),('Caption',9,GRAY,False)]:
    st=styles[nm]; st.font.name='Aptos'; st.font.size=Pt(size); st.font.bold=bold; st.font.color.rgb=RGBColor.from_string(color.replace('#',''))
styles['Normal'].paragraph_format.space_after=Pt(6); styles['Normal'].paragraph_format.line_spacing=1.12
for nm in ['Heading 1','Heading 2','Heading 3']: styles[nm].paragraph_format.space_before=Pt(12); styles[nm].paragraph_format.space_after=Pt(5); styles[nm].paragraph_format.keep_with_next=True

# Header/footer
hp=sec.header.paragraphs[0]; hp.text="EVIDENCE-GROUNDED AI FOR LEGACY DATABASE RCA  |  INTERNAL DRAFT v0.2"; hp.alignment=WD_ALIGN_PARAGRAPH.CENTER; hp.runs[0].font.size=Pt(8); hp.runs[0].font.color.rgb=RGBColor.from_string(GRAY.replace('#',''))
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; fp.add_run("Mukesh Dabi  •  "); add_field(fp.add_run(),"PAGE")

p=doc.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("Design and Evaluation of an Evidence-Grounded Agentic AI Framework for Root-Cause Analysis in Legacy Database Systems")
p=doc.add_paragraph(style='Subtitle'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("Verified Research Draft v0.2")
for line in ["Author: Mukesh Dabi","Affiliation: [Institution or independent-research affiliation]","Corresponding email: [Email]","Document date: 19 July 2026","Status: Internal draft — not ready for submission"]:
    q=doc.add_paragraph(line); q.alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(); q=doc.add_paragraph("Reproducibility anchor",style='Heading 2'); q.alignment=WD_ALIGN_PARAGRAPH.CENTER
for line in [f"Frozen commit: {summary['provenance']['commit']}",f"Release tag: {summary['provenance']['tag']}",f"Manifest SHA-256: {summary['provenance']['manifest_sha256']}",f"Run: {summary['provenance']['run_name']}"]:
    q=doc.add_paragraph(line); q.alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()
doc.add_heading('Abstract',level=1)
doc.add_paragraph("Operational support for legacy relational systems requires schema discovery, safe diagnostic querying, correlation of records, and defensible separation of symptoms from causes. This paper presents an implemented evidence-grounded agentic framework that stages intent analysis, entity extraction, metadata discovery, object and relationship ranking, read-only SQL planning and validation, evidence collection, an evidence gate, structured reasoning, and report composition. The frozen evaluation contains 125 reproducible scenarios across five domains, with five SQL lifecycle scripts per scenario (625 scripts). All 125 scenarios started and reached a terminal state; cleanup passed for 125/125. Strict end-to-end success was 46/125 = 36.8% (95% Wilson CI 28.9%–45.5%). Conditional application accuracy was 48/51 = 94.12% (95% Wilson CI 84.1%–98.0%) among usable answers eligible for deterministic validation. These denominators measure different properties. Fifty-five of 125 runs were classified as non-timeout provider failures recorded only as HTTPError, and 19 lacked sufficient evidence. Because HTTP status was not retained, rate limiting is plausible but not established. The results support the feasibility of evidence-grounded database investigation when the reasoning path completes, but they do not establish production readiness, superiority to baselines, or generalization to live enterprise databases.")
doc.add_paragraph("Keywords: agentic AI; legacy databases; root-cause analysis; Text-to-SQL; evidence grounding; SQL safety; AI evaluation; provenance")
doc.add_heading('Table of Contents',level=1); p=doc.add_paragraph(); add_field(p.add_run(),"TOC \\o \"1-3\" \\h \\z \\u")

doc.add_heading('1. Introduction',level=1)
doc.add_paragraph("Legacy applications often encode business behavior across schemas, stored procedures, status tables, queues, and operational conventions. A fluent answer is inadequate when it cannot be traced to executed read-only queries and observed records. The research problem is therefore to produce useful root-cause explanations over unfamiliar databases while enforcing action safety, preserving provenance, and measuring both answer quality and operational reliability.")
doc.add_paragraph("This work contributes an implemented staged architecture; deterministic SQL and evidence controls; a 125-scenario synthetic, cross-domain benchmark; and an evaluation protocol that distinguishes conditional application accuracy from strict end-to-end reliability. It does not claim algorithmic novelty over all database agents, production safety, or superiority to alternative systems because no baseline or human study was executed.")
doc.add_heading('1.1 Research questions and hypotheses',level=2)
for x in ["RQ1: What conditional root-cause accuracy is achieved among usable, eligible answers?","RQ2: What strict success rate is achieved when all requested scenarios and failure modes remain in the denominator?","RQ3: How do outcomes vary by domain and difficulty?","RQ4: Which evidence, provider, application, and judge failures limit reliability?","RQ5: Which claims can be verified from persisted evidence and deterministic checks?"]:
    doc.add_paragraph(x,style='List Bullet')
doc.add_paragraph("The draft treats the expected gap between conditional accuracy and strict reliability as a descriptive hypothesis. Significance tests against a pre-registered baseline are not reported because no baseline threshold or comparator experiment was frozen before analysis.")

doc.add_heading('2. Related Work',level=1)
doc.add_paragraph("Spider formalized cross-domain Text-to-SQL generalization over unseen schemas [1]. Spider 2.0 extends the problem toward enterprise workflows [2]. ReAct demonstrated interleaving reasoning and external action [3]. These strands motivate tool-using database assistants but do not by themselves guarantee query safety or evidential support. LLM-as-a-judge studies also warn that model judgments can diverge from human judgments and exhibit prompt and leniency effects [4]. Accordingly, the present evaluation places deterministic verification before secondary semantic judging and identifies human validation as missing work.")

doc.add_heading('3. System Architecture',level=1)
doc.add_paragraph("The implementation uses deterministic orchestration around optional LLM reasoning. Figure 1 maps repository components to the deployed investigation path. The model does not receive a database connection; it receives a masked evidence package and its output is retained only when linked to evidence.")
add_figure(doc,1,figs[0],"Overall system architecture and trust boundaries.")
add_table(doc,"Table 1. Verified system components and responsibilities",["Component","Verified responsibility","Evidence source"],[
 ["Intent agent","Classifies request intent before discovery and planning","agents/intent_agent.py"],["Entity extraction/resolution","Preserves identifiers and links canonical entities to evidence","entity_extraction_agent.py; entity_resolution_service.py"],["Metadata discovery","Extracts and caches tables, views, columns, indexes, keys and procedures","db/connector.py; metadata_search_service.py"],["Object/relationship ranking","Ranks objects using lexical, index, foreign-key, and intent signals","object_ranking_agent.py; relationship_analysis_agent.py"],["SQL planning/validation","Plans diagnostic queries and enforces a read-only statement allowlist","safe_sql_planner.py; sql_validator.py"],["Evidence gate","Checks entity, affected rows, condition, and required relationships","combined_evidence_gate.py"],["Reasoning","Optional OpenAI reasoning over masked evidence package","llm_reasoning_agent.py; config.py"],["Report composer","Produces structured report content and downloadable formats","report_composer_agent.py"],["Evaluation","Prepares fixtures, invokes public API, verifies, judges, persists, cleans","evaluation/**"],
], [1.25,3.55,1.7])
doc.add_heading('3.1 Investigation sequence',level=2); doc.add_paragraph("Figure 2 shows the implemented order. Metadata and evidence constrain later stages; rejected SQL is not executed."); add_figure(doc,2,figs[1],"Investigation sequence from natural-language question to report.")

doc.add_heading('4. Evidence and Safety Controls',level=1)
doc.add_paragraph("The implementation allows SELECT, SHOW, DESCRIBE, DESC, and EXPLAIN statement families and rejects writes and stored-procedure execution in the evidence path. The default maximum investigation result is 100 rows, and full-table scans are disabled by default. Evaluation connections are documented as separate contained read-only identities with db_datareader and explicit denials for mutation and procedure execution. These controls reduce risk but are not a formal proof of SQL safety and do not establish fitness for production use.")
add_table(doc,"Table 2. Safety controls",["Control","Implemented behavior","Qualification"],[
 ["Statement allowlist","SELECT, SHOW, DESCRIBE, DESC, EXPLAIN","Not a claim of complete semantic SQL parsing"],["Result bound","MAX_INVESTIGATION_ROWS defaults to 100","Configuration dependent"],["Full scans","Disabled by default","May be explicitly configured"],["Identity isolation","Separate read-only evaluation identity","Documented deployment control"],["Approval boundary","Suggested verification queries require user approval","Applies to verification flow"],["Audit logging","Feature enabled by default","Coverage depends on execution path and deployment"],
],[1.45,2.8,2.25])
doc.add_heading('4.1 Evidence-gate behavior',level=2); doc.add_paragraph("The gate determines whether connected evidence supports reasoning. Figure 3 presents the reporting semantics: supported claims may be confirmed; partial support requires qualification; insufficient evidence withholds a root cause; contradiction rejects the claim. The frozen benchmark directly measures gate pass/rejection, not a complete four-class calibration study."); add_figure(doc,3,figs[2],"Evidence-gate decisions and permitted report language.")

doc.add_heading('5. Benchmark and Method',level=1)
doc.add_paragraph("The frozen inventory contains 125 active scenarios: 25 per domain. Each scenario supplies scenario.json plus baseline_reset.sql, precondition.sql, inject.sql, verify.sql, and cleanup.sql. This yields 625 SQL scripts. The run used MySQL databases on 127.0.0.1, application version 0.1.0, gpt-4.1-mini for both reasoning and judging, prompt versions evidence-grounded-v1 and ai-judge-v2-entity-provenance, concurrency 1, and a 600-second scenario timeout.")
add_figure(doc,4,figs[3],"Benchmark composition by domain and difficulty.")
add_table(doc,"Table 3. Benchmark distribution by domain",["Domain","Scenarios","Share"],[[k,v['requested'],f"{v['requested']/125:.1%}"] for k,v in summary['by_domain'].items()],[2.8,1.7,2.0])
add_table(doc,"Table 4. Benchmark distribution by difficulty",["Difficulty","Scenarios","Share"],[[k,v['requested'],f"{v['requested']/125:.1%}"] for k,v in summary['by_difficulty'].items()],[2.8,1.7,2.0])
add_table(doc,"Table 5. Experimental environment",["Field","Verified value"],[["Frozen commit",summary['provenance']['commit']],["Release tag",summary['provenance']['tag']],["Manifest SHA-256",summary['provenance']['manifest_sha256']],["Run name",summary['provenance']['run_name']],["Concurrency","1"],["Scenario timeout","600 seconds"],["Reasoning / judge model","gpt-4.1-mini / gpt-4.1-mini"],["Database engine / host","MySQL / 127.0.0.1"],["Requested / started / terminal / incomplete","125 / 125 / 125 / 0"],["Cleanup / post-run fixture audit","125/125 / 125/125 valid"]],[2.2,4.3])
doc.add_heading('5.1 Evaluation definitions',level=2)
doc.add_paragraph("Strict pass uses all requested scenarios: PASS/125. Conditional application accuracy uses deterministic passes among usable AI answers eligible for deterministic validation: 48/51. Provider failures, evidence-insufficient cases, and judge failures remain in strict reliability accounting but are not all eligible for conditional accuracy. AI Judge scores are secondary semantic measurements, not ground truth.")
add_table(doc,"Table 6. AI Judge dimensions",["Dimension","Weight"],[["Root-cause correctness","30%"],["Evidence correctness","25%"],["Database-object discovery","10%"],["Fix correctness","10%"],["Citation correctness","10%"],["Safety","10%"],["Completeness","5%"]],[4.6,1.9])

doc.add_heading('6. Results',level=1)
doc.add_paragraph("The strict pass rate was 46/125 = 36.8% (95% Wilson CI 28.9%–45.5%). Conditional application accuracy was 48/51 = 94.12% (95% Wilson CI 84.1%–98.0%). The usable-answer rate was 51/125 = 40.8%. Figure 5 deliberately uses distinct denominators.")
add_figure(doc,5,figs[4],"Application accuracy, strict pass rate, and provider failures with denominators.")
add_table(doc,"Table 7. Overall results",["Measure","Numerator","Denominator","Value"],[["Strict end-to-end pass",46,125,"36.8%"],["Conditional application accuracy",48,51,"94.12%"],["Usable AI-answer rate",51,125,"40.8%"],["Evidence-gate pass",106,125,"84.8% of all scenarios"],["Evidence-gate rejection",14,125,"11.2% of all scenarios"],["Other evidence-insufficient",5,125,"4.0% of all scenarios"],["Judge completion",49,51,"96.08%"],["Cleanup",125,125,"100%"]],[2.9,1.1,1.2,1.4])
doc.add_heading('6.1 Failure taxonomy',level=2); doc.add_paragraph("Fifty-five provider failures dominate the strict outcome. They are recorded as HTTPError and PROVIDER_OTHER_FAILURE, not timeouts. Because response status was not retained, a specific cause such as rate limiting cannot be verified. Nineteen scenarios were evidence-insufficient, three were application-incorrect, and two had judge failures."); add_figure(doc,6,figs[5],"Counts for every frozen primary classification.")
add_table(doc,"Table 8. Failure taxonomy",["Classification","Count","Interpretation"],[[k,v,"Primary frozen classification"] for k,v in summary['classifications'].items()],[2.7,1.0,2.8])
doc.add_heading('6.2 Domain and difficulty results',level=2); add_figure(doc,7,figs[6],"Strict passes by domain; each domain denominator is 25.")
domain_rows=[]
for k,v in summary['by_domain'].items():
    acc="Not estimable (0/0)" if not v['eligible'] else f"{v['pass']+v.get('application_incorrect',0) if False else (round(v['application_accuracy']*v['eligible']))}/{v['eligible']} = {v['application_accuracy']:.2%}"
    domain_rows.append([k,f"{v['pass']}/25 = {v['exact_pass_rate']:.1%}",v['eligible'],acc,v['provider_other_failure'],v['evidence_insufficient']])
add_table(doc,"Table 9. Domain-level outcomes",["Domain","Strict pass","Eligible","Conditional accuracy","Provider failures","Evidence insufficient"],domain_rows,[1.0,1.25,.7,1.7,.9,.95])
diff_rows=[]
for k,v in summary['by_difficulty'].items(): diff_rows.append([k,f"{v['pass']}/{v['requested']} = {v['exact_pass_rate']:.1%}",v['eligible'],f"{round(v['application_accuracy']*v['eligible'])}/{v['eligible']} = {v['application_accuracy']:.2%}",v['provider_other_failure'],v['evidence_insufficient']])
add_table(doc,"Table 10. Difficulty-level outcomes",["Difficulty","Strict pass","Eligible","Conditional accuracy","Provider failures","Evidence insufficient"],diff_rows,[1.0,1.25,.7,1.7,.9,.95])

doc.add_heading('7. End-to-End Evaluation Architecture',level=1); doc.add_paragraph("The runner loads a frozen contract, verifies the target, applies preconditions and injected defects, invokes the application through POST /chat/ask, polls the saved investigation, captures persisted evidence, applies deterministic verification and secondary judging, persists results, and runs cleanup. Figure 8 separates application execution from evaluator logic."); add_figure(doc,8,figs[7],"Scenario lifecycle and evaluation boundaries.")

doc.add_heading('8. Discussion',level=1)
doc.add_paragraph("The 94.12% conditional accuracy indicates that most usable eligible outputs satisfied deterministic expectations. It must not be read as end-to-end system accuracy: only 51/125 runs entered that denominator. The 36.8% strict rate better represents the observed release-candidate workflow under the recorded provider conditions. Domain comparisons are heavily confounded by execution order because provider failures cluster in later orders, banking, and shipping runs. Banking and shipping conditional accuracy is not estimable, rather than zero, because neither domain produced eligible responses.")
doc.add_paragraph("Evidence gating prevented at least 14 unsupported reasoning invocations, while five additional cases were insufficient for other recorded reasons. This is operational evidence of abstention behavior, not proof that every unsupported claim is detected. Likewise, the AI Judge’s mean score of 80.308 across 49 completed judgments should be treated as an auxiliary indicator because no human agreement study was performed.")

doc.add_heading('9. Threats to Validity and Limitations',level=1)
add_table(doc,"Table 11. Threats to validity",["Threat","Observed issue","Mitigation / required work"],[
 ["Construct validity","Conditional accuracy excludes 74 scenarios","Always report numerator, denominator, and eligibility rule"],["Internal validity","55 provider failures cluster by execution period/domain","Repeat randomized or blocked runs with persisted HTTP status"],["External validity","Synthetic five-domain fixtures","Validate on de-identified real incidents and additional engines"],["Evaluator validity","AI Judge lacks human calibration","Conduct blinded expert review and inter-rater analysis"],["Statistical validity","No pre-registered baseline or powered comparison","Freeze hypotheses, baseline, and analysis plan before rerun"],["Safety validity","Read-only controls are not formal verification","Adversarial SQL testing, parser review, least-privilege deployment audit"],["Reproducibility","Current full pytest collection hit two permission errors","Resolve environment permissions and publish a clean test artifact"],
],[1.35,2.45,2.7])
doc.add_paragraph("Additional limitations include zero recorded cost that is not a billing assertion; no latency service-level objective; no human evaluation; no baseline model; no ablation of evidence gating; no causal attribution for provider HTTPError; and no evidence that the system is production-safe.")

doc.add_heading('10. Conclusion and Future Work',level=1)
doc.add_paragraph("The verified implementation demonstrates a practical evidence-grounded investigation pipeline with explicit SQL controls, provenance-aware evidence, abstention, reporting, and reproducible synthetic evaluation. The frozen benchmark establishes high conditional correctness among usable eligible outputs (48/51) but low observed end-to-end reliability (46/125), primarily because 55 provider HTTPError failures and 19 evidence-insufficient outcomes remain in the strict denominator. The system is therefore a research prototype requiring further reliability, safety, and evaluation work.")
doc.add_paragraph("Priority future work is to persist provider status and request identifiers; randomize or block scenario order; rerun the full benchmark under stable provider conditions; add deterministic and non-agentic baselines; conduct blinded database-expert evaluation; calibrate the AI Judge; test PostgreSQL and SQL Server end-to-end; perform adversarial SQL-safety testing; and pre-register statistical hypotheses.")

doc.add_heading('References',level=1)
refs=[
"[1] Yu, T., et al. (2018). Spider: A large-scale human-labeled dataset for complex and cross-domain semantic parsing and Text-to-SQL task. EMNLP, 3911–3921. https://doi.org/10.18653/v1/D18-1425",
"[2] Lei, F., et al. (2024). Spider 2.0: Evaluating language models on real-world enterprise Text-to-SQL workflows. arXiv. https://arxiv.org/abs/2411.07763",
"[3] Yao, S., et al. (2023). ReAct: Synergizing reasoning and acting in language models. ICLR. https://openreview.net/forum?id=WE_vluYUL-X",
"[4] Thakur, A. S., et al. (2024). Judging the judges: Evaluating alignment and vulnerabilities in LLMs-as-judges. arXiv. https://arxiv.org/abs/2406.12624",
"[5] LegacyDB Support Copilot. (2026). README, evaluation documentation, source code, and frozen benchmark artifacts at commit d5815fd509a13cb9dd3eec28c859c79f205d3c80. Internal repository artifact.",
]
for r in refs: doc.add_paragraph(r)

doc.add_heading('Appendix A. Benchmark Provenance',level=1)
add_table(doc,"Table A1. Frozen artifact identifiers",["Artifact","Value"],[["Commit",summary['provenance']['commit']],["Tag",summary['provenance']['tag']],["Manifest SHA-256",summary['provenance']['manifest_sha256']],["Summary checksum","8d14d69e55700a847b1b92844632d98cd30692a047fc61cba3852492b52e7dcc"],["Results CSV checksum","8709b1bf60025b6274bda24383606d56b6c04e5d64f9f46db5b409f817fdb4ca"]],[2.0,4.5])
doc.add_heading('Appendix B. Reproducibility Checklist',level=1)
for x in ["Frozen scenario list and manifest checksum recorded.","Preflight passed; requested/started/terminal = 125/125/125.","Cleanup passed for 125/125; fixture audit after run was 125/125 valid.","Reasoning and judge model and prompt versions recorded.","Application and evaluator commit matched the release tag.","Known provider-status logging limitation retained.","No baseline or human-evaluation result claimed."]:
    doc.add_paragraph(x,style='List Bullet')
doc.add_heading('Appendix C. Missing Experiments',level=1)
for x in ["Randomized full rerun with stable provider access and persisted HTTP status.","Non-agentic, Text-to-SQL-only, and deterministic baselines.","Evidence-gate ablation and counterfactual contradiction tests.","Blinded expert evaluation with agreement statistics.","Cross-engine replication on PostgreSQL and SQL Server.","Adversarial SQL safety and performance-load evaluation."]:
    doc.add_paragraph(x,style='List Bullet')
doc.add_heading('Appendix D. Artifact Inventory',level=1)
doc.add_paragraph("Primary evidence reviewed: application agents and services; database connectors; authentication, tenancy, and configuration code; evaluation framework, runners, deterministic validators, AI Judge, and reports; 125 scenario contracts and 625 SQL scripts; frozen provenance, checksums, summary JSON, results CSV, release report, preflight log, fixture audits, and selected execution logs; README, BENCHMARK.md, deployment documentation, migrations, and tests.")
doc.add_heading('Appendix E. Verification and Correction Log',level=1)
add_table(doc,"Table E1. Important corrections",["Section","Existing statement","Actual evidence","Correction","Source"],discrepancies,[.8,1.35,1.65,1.6,1.1])

docx=OUT/"Evidence_Grounded_AI_Research_Paper_Draft_v0.2_Verified.docx"; doc.save(docx)

report=f"""# Research Paper Verification Report

## Outcome

The draft was rebuilt as verified version 0.2 using the frozen release benchmark and the implemented repository. Submission-readiness status: **not ready for submission**. The paper is technically reviewable, but a stable-provider rerun, baseline experiments, and human evaluation are still missing.

## Evidence reviewed

- Frozen benchmark run `{summary['provenance']['run_name']}` at commit `{summary['provenance']['commit']}` and tag `{summary['provenance']['tag']}`.
- Manifest SHA-256 `{summary['provenance']['manifest_sha256']}` and artifact checksums.
- 125 scenario contracts and 625 lifecycle SQL scripts.
- Application agents, metadata/relationship services, SQL safety controls, evidence gate, LLM integration, report composition, authentication/tenancy, audit configuration, evaluation runner, deterministic validator, AI Judge, persistence, documentation, and selected tests/logs.

## Claims verified or corrected

- Strict pass: **46/125 = 36.8%** (95% Wilson CI 28.9%–45.5%).
- Conditional application accuracy: **48/51 = 94.12%** (95% Wilson CI 84.1%–98.0%).
- Provider failures: **55/125**, recorded as `HTTPError` / `PROVIDER_OTHER_FAILURE`; HTTP status was not retained.
- Evidence insufficient: **19/125**; application incorrect: **3/125**; judge failure: **2/125**.
- Requested/started/terminal/incomplete: **125/125/125/0**; cleanup and post-run fixture audit: **125/125**.
- Domain distribution: 25 each across payroll, clinic, orders, banking, and shipping.
- Difficulty distribution: medium 50, hard 45, easy 20, expert 10.
- AI Judge rubric: 30/25/10/10/10/10/5 over root cause, evidence, object discovery, fix, citation, safety, and completeness.

## Not verifiable / explicitly withheld

- The provider failure was caused by rate limiting (plausible, not proven).
- Production safety or production readiness.
- Superiority over baselines.
- Human-level quality or human agreement.
- Statistical significance against a pre-registered comparator.
- Billing cost (recorded zero is not a billing assertion).

## Test-status caveat

A fresh `pytest --collect-only -q` on 2026-07-19 enumerated more than 1,000 cases but was interrupted by `PermissionError` reading `.env.evaluation` in two evaluation modules. The paper therefore does not claim a clean current full-suite run. Frozen benchmark preflight, completion, cleanup, and fixture-audit evidence are reported separately.

## Missing experiments and future work

Stable-provider randomized rerun; persisted HTTP status/request IDs; baseline and ablation studies; blinded expert evaluation; AI Judge calibration; cross-engine replication; adversarial SQL-safety testing; and a pre-registered statistical plan.
"""
(OUT/"Research_Paper_Verification_Report.md").write_text(report,encoding="utf-8")

# Workbook (artifact-tool runtime was unavailable in this session; deterministic OOXML workbook fallback).
wb=Workbook(); ws=wb.active; ws.title="Correction Log"
headers=["Section","Existing statement","Actual evidence","Required correction","Source file or report"]
ws.append(headers)
for row in discrepancies: ws.append(row)
ws.freeze_panes="A2"; ws.auto_filter.ref=f"A1:E{ws.max_row}"
tbl=Table(displayName="CorrectionLog",ref=f"A1:E{ws.max_row}"); tbl.tableStyleInfo=TableStyleInfo(name="TableStyleMedium2",showRowStripes=True,showColumnStripes=False); ws.add_table(tbl)
for c in ws[1]: c.font=Font(bold=True,color="FFFFFF"); c.fill=PatternFill('solid',fgColor=NAVY.replace('#','')); c.alignment=Alignment(wrap_text=True,vertical='center')
widths=[22,42,48,48,42]
for i,w in enumerate(widths,1): ws.column_dimensions[chr(64+i)].width=w
for row in ws.iter_rows(min_row=2):
    for c in row: c.alignment=Alignment(wrap_text=True,vertical='top')
    ws.row_dimensions[c.row].height=65
src=wb.create_sheet("Evidence Sources"); src.append(["Artifact","Verified value / purpose"])
for row in [["Frozen commit",summary['provenance']['commit']],["Release tag",summary['provenance']['tag']],["Manifest SHA-256",summary['provenance']['manifest_sha256']],["Benchmark summary","Counts, denominators, domain and difficulty breakdown"],["Repository source","Architecture, safety, auth, evidence, evaluation implementation"]]: src.append(row)
for c in src[1]: c.font=Font(bold=True,color="FFFFFF"); c.fill=PatternFill('solid',fgColor=NAVY.replace('#',''))
src.column_dimensions['A'].width=28; src.column_dimensions['B'].width=90; src.freeze_panes='A2'
for row in src.iter_rows():
    for c in row: c.alignment=Alignment(wrap_text=True,vertical='top')
wb.save(OUT/"Research_Paper_Discrepancy_Log.xlsx")

print(docx)
