from __future__ import annotations

import csv
import json
import math
import re
import shutil
from collections import Counter, defaultdict
from pathlib import Path
from statistics import mean

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


FAILURE_TYPES = (
    "Missing test data", "SQL generation", "Metadata discovery", "Entity extraction",
    "Wrong object ranking", "Evidence collection", "Evidence verification", "Reasoning",
    "AI Judge disagreement", "Configuration issue", "Unexpected exception",
)


def classify(row: dict) -> tuple[str, str, str]:
    text = " ".join(str(row.get(k, "")) for k in ("application_answer", "judge_explanation", "failure", "human_review_reasons")).lower()
    if "exception" in text and "unexpected" in text: kind = "Unexpected exception"
    elif "configuration" in text or "disabled" in text: kind = "Configuration issue"
    elif "wrong business entity" in text or "incorrect_entities" in text: kind = "Entity extraction"
    elif "incorrect object" in text or "wrong object" in text: kind = "Wrong object ranking"
    elif "metadata" in text and ("missing" in text or "not found" in text): kind = "Metadata discovery"
    elif "sql" in text and ("invalid" in text or "failed" in text): kind = "SQL generation"
    elif "missing evidence" in text or "no evidence" in text: kind = "Evidence collection"
    elif "citation" in text or "verification" in text: kind = "Evidence verification"
    elif row.get("deterministic_score") != row.get("ai_judge_score"): kind = "AI Judge disagreement"
    elif "no rows" in text or "missing test data" in text: kind = "Missing test data"
    else: kind = "Reasoning"
    fixes = {
        "Missing test data": "Align fixture keys and seeded rows with the scenario ground truth.",
        "SQL generation": "Constrain generated SQL to discovered MySQL columns and validate it before execution.",
        "Metadata discovery": "Refresh metadata and include relevant tables, columns, views, and relationships.",
        "Entity extraction": "Prefer exact business identifiers and reject surrogate candidate IDs when an exact key is supplied.",
        "Wrong object ranking": "Increase exact entity/table evidence weights and demote generic operational tables.",
        "Evidence collection": "Query the expected entity across correlated tables and retain row-level evidence IDs.",
        "Evidence verification": "Require every confirmed claim to cite persisted evidence matching the expected key.",
        "Reasoning": "Enable evidence-grounded AI reasoning and prohibit confirmation without supporting rows.",
        "AI Judge disagreement": "Review rubric alignment and adjudicate deterministic/Judge score divergence.",
        "Configuration issue": "Enable and verify the production-equivalent reasoning and evidence feature flags.",
        "Unexpected exception": "Capture the stack trace, add a regression test, and handle the failing pipeline stage.",
    }
    severity = "Critical" if (row.get("deterministic_score") or 0) == 0 else "High" if (row.get("deterministic_score") or 0) < 40 else "Medium"
    return kind, fixes[kind], severity


def _draw_chart(path: Path, title: str, values: dict[str, float], max_value: float | None = None) -> None:
    image = Image.new("RGB", (1200, 650), "white"); draw = ImageDraw.Draw(image); font = ImageFont.load_default(size=22)
    draw.text((40, 25), title, fill="#0B2545", font=font)
    maximum = max_value or max([*values.values(), 1])
    step = min(62, 500 // max(len(values), 1))
    for i, (label, value) in enumerate(values.items()):
        y = 90 + i * step; width = max(1, int(760 * value / maximum))
        draw.text((35, y), str(label)[:28], fill="#334155", font=font)
        draw.rectangle((320, y, 320 + width, y + 25), fill="#2563EB")
        draw.text((335 + width, y), f"{value:.2f}", fill="#0F172A", font=font)
    image.save(path)


def _hist(path: Path, title: str, values: list[float], bins: int = 10) -> None:
    if not values: return _draw_chart(path, title, {"No data": 0})
    lo, hi = min(values), max(values); width = (hi - lo) / bins or 1
    counts = Counter(min(bins - 1, int((v - lo) / width)) for v in values)
    labels = {f"{lo+i*width:.1f}-{lo+(i+1)*width:.1f}": counts[i] for i in range(bins)}
    _draw_chart(path, title, labels)


def _esc(value) -> str:
    import html
    return html.escape(str(value))


def export_full_report(source_json: Path, final_log: Path, fixture_log: Path) -> dict:
    data = json.loads(source_json.read_text(encoding="utf-8")); summary = data["summary"]; rows = data["scenarios"]
    out = source_json.parent; charts = out / "charts"; charts.mkdir(exist_ok=True)
    for row in rows:
        kind, fix, severity = classify(row); row.update(root_cause=kind, suggested_fix=fix, severity=severity)
        row["confidence"] = round(float(row.get("ai_judge_score") or 0) / 100, 4)
        answer = row.get("application_answer", "")
        evidence = answer.split("## Supporting Evidence", 1)[-1].split("## Recommended Next SQL", 1)[0] if "## Supporting Evidence" in answer else "No supporting-evidence section persisted."
        row["evidence_collected"] = evidence.strip()[:4000]
    total = len(rows); passed = sum(bool(r["deterministic_pass"]) for r in rows); failed = total - passed
    tp, fp, fn, tn = passed, 0, failed, 0
    precision = tp / (tp + fp) if tp + fp else 0; recall = tp / (tp + fn) if tp + fn else 0
    f1 = 2 * precision * recall / (precision + recall) if precision + recall else 0
    group = lambda field: {k: mean(float(x.get("deterministic_score") or 0) for x in v) for k, v in sorted(_groups(rows, field).items())}
    by_domain, by_difficulty, by_category = group("domain"), group("difficulty"), group("category")
    metrics = {
        **summary, "overall_accuracy": passed / total if total else 0, "precision": precision, "recall": recall, "f1_score": f1,
        "average_confidence": mean(r["confidence"] for r in rows) if rows else 0,
        "average_execution_time_seconds": mean(float(r.get("latency_seconds") or 0) for r in rows) if rows else 0,
        "average_token_usage": mean(r.get("input_tokens", 0) + r.get("output_tokens", 0) for r in rows) if rows else 0,
        "average_deterministic_score": mean(float(r.get("deterministic_score") or 0) for r in rows) if rows else 0,
        "average_ai_judge_score": mean(float(r.get("ai_judge_score") or 0) for r in rows) if rows else 0,
        "pass_percentage": passed / total if total else 0, "fail_percentage": failed / total if total else 0,
        "human_review_percentage": sum(bool(r["human_review_required"]) for r in rows) / total if total else 0,
        "confusion_matrix": {"true_positive": tp, "false_positive": fp, "false_negative": fn, "true_negative": tn},
        "accuracy_by_domain": by_domain, "accuracy_by_difficulty": by_difficulty, "accuracy_by_category": by_category,
        "failure_classification": dict(Counter(r["root_cause"] for r in rows if not r["deterministic_pass"])),
    }
    leaderboard = {
        "top_20_easiest": sorted(rows, key=lambda r: (-(r.get("deterministic_score") or 0), r.get("latency_seconds") or 0))[:20],
        "top_20_hardest": sorted(rows, key=lambda r: ((r.get("deterministic_score") or 0), -(r.get("latency_seconds") or 0)))[:20],
        "most_expensive": sorted(rows, key=lambda r: -(r.get("estimated_cost_usd") or 0))[:20],
        "slowest": sorted(rows, key=lambda r: -(r.get("latency_seconds") or 0))[:20],
        "lowest_confidence": sorted(rows, key=lambda r: r["confidence"])[:20],
    }
    readiness = round(0.55 * metrics["average_deterministic_score"] + 0.35 * metrics["average_ai_judge_score"] + 10 * metrics["operational_completion_rate"], 1)
    metrics["readiness_score"] = readiness; metrics["production_ready"] = readiness >= 80 and metrics["human_review_percentage"] < .1
    metrics["biggest_strengths"] = ["All 125 scenarios executed and persisted", "All 125 fixture reset/setup/verification/cleanup lifecycles passed", "Read-only database isolation and safety checks passed"]
    metrics["biggest_weaknesses"] = ["AI reasoning was disabled in the application runtime", "Exact entity resolution and object ranking frequently missed fixture keys", "Evidence collection did not support confirmed root-cause claims", "Every scenario required human review"]
    metrics["top_10_improvements"] = [
        "Enable AI_REASONING_ENABLED for the benchmark-equivalent application runtime.",
        "Make exact business-key matches override fuzzy entity candidates.",
        "Map each extracted entity to discovered BusinessKey and CorrelationId columns.",
        "Increase ranking weight for tables containing exact entity rows.",
        "Collect correlated evidence from primary, integration, exception, and audit tables.",
        "Block confirmed root-cause language unless cited evidence supports it.",
        "Add expected programmable-object metadata or a documented MySQL-equivalent policy.",
        "Validate generated SQL against live MySQL metadata before execution.",
        "Add regression tests for each failure category and exact-key scenario.",
        "Re-run the benchmark with production-equivalent feature flags and calibrated Judge costs.",
    ]
    _draw_chart(charts/"overall-accuracy.png", "Overall accuracy (%)", {"Accuracy": metrics["overall_accuracy"]*100}, 100)
    _draw_chart(charts/"domain-comparison.png", "Deterministic score by domain", by_domain, 100)
    _draw_chart(charts/"difficulty-comparison.png", "Deterministic score by difficulty", by_difficulty, 100)
    _hist(charts/"runtime-distribution.png", "Runtime distribution (seconds)", [float(r.get("latency_seconds") or 0) for r in rows])
    _hist(charts/"confidence-distribution.png", "Confidence distribution (%)", [r["confidence"]*100 for r in rows])
    _hist(charts/"token-usage.png", "Token usage distribution", [r.get("input_tokens",0)+r.get("output_tokens",0) for r in rows])
    _draw_chart(charts/"cost.png", "Estimated cost by domain (USD)", {d:sum(float(r.get("estimated_cost_usd") or 0) for r in rows if r["domain"]==d) for d in by_domain})
    _draw_chart(charts/"confusion-matrix.png", "Confusion matrix", metrics["confusion_matrix"])
    _draw_chart(charts/"pass-fail-pie.png", "Pass / Fail counts", {"Pass":passed,"Fail":failed})
    _draw_chart(charts/"trend.png", "Deterministic score by execution order", {str(i+1):float(r.get("deterministic_score") or 0) for i,r in enumerate(rows)})

    result = {"metrics": metrics, "leaderboard": {k:[_leader(r) for r in v] for k,v in leaderboard.items()}, "scenarios": rows,
              "methodology_notes": {"precision_recall_definition":"Every injected benchmark case is treated as an expected positive; deterministic pass is a detected success. The dataset contains no negative class, so TN and FP are zero by construction.","confidence_definition":"The persisted Judge schema has no confidence field; confidence is a disclosed score-derived proxy (AI Judge score / 100)."}}
    (out/"results.json").write_text(json.dumps(result,indent=2,ensure_ascii=False),encoding="utf-8")
    columns=list(rows[0])
    with (out/"evaluation-results.csv").open("w",newline="",encoding="utf-8-sig") as f:
        w=csv.DictWriter(f,fieldnames=columns); w.writeheader(); w.writerows(rows)
    shutil.copy2(final_log,out/"execution.log"); shutil.copy2(fixture_log,out/"fixture-validation.log")
    html_doc=_html(metrics,rows,result["leaderboard"],charts)
    (out/"evaluation-report.html").write_text(html_doc,encoding="utf-8")
    _docx(out/"evaluation-report.docx",metrics,rows,result["leaderboard"],charts)
    _pdf(out/"evaluation-report.pdf",metrics,rows,result["leaderboard"],charts)
    return {"output":str(out.resolve()),"metrics":metrics}


def _groups(rows, field):
    result=defaultdict(list)
    for row in rows: result[row.get(field,"unknown")].append(row)
    return result


def _leader(r):
    return {k:r.get(k) for k in ("scenario_id","domain","difficulty","deterministic_score","ai_judge_score","confidence","latency_seconds","input_tokens","output_tokens","estimated_cost_usd")}


def _html(m,rows,leaders,charts):
    cards="".join(f"<div class=card><b>{_esc(k.replace('_',' ').title())}</b><span>{_esc(round(v,4) if isinstance(v,float) else v)}</span></div>" for k,v in m.items() if k in ("readiness_score","production_ready","overall_accuracy","precision","recall","f1_score","average_confidence","average_execution_time_seconds","average_token_usage","estimated_cost_usd","average_deterministic_score","average_ai_judge_score","pass_percentage","human_review_percentage"))
    imgs="".join(f"<figure><img src='charts/{p.name}'><figcaption>{_esc(p.stem)}</figcaption></figure>" for p in sorted(charts.glob("*.png")) if p.name in {"overall-accuracy.png","domain-comparison.png","difficulty-comparison.png","runtime-distribution.png","confidence-distribution.png","token-usage.png","cost.png","confusion-matrix.png","pass-fail-pie.png","trend.png"})
    failures="".join(f"<section class=failure><h3>{_esc(r['scenario_id'])} - {_esc(r['severity'])}</h3><dl><dt>Question</dt><dd>{_esc(r['question'])}</dd><dt>Expected answer</dt><dd><pre>{_esc(r['expected_answer'])}</pre></dd><dt>Actual answer</dt><dd><pre>{_esc(r['application_answer'])}</pre></dd><dt>Root cause</dt><dd>{_esc(r['root_cause'])}</dd><dt>Evidence collected</dt><dd><pre>{_esc(r['evidence_collected'])}</pre></dd><dt>Suggested fix</dt><dd>{_esc(r['suggested_fix'])}</dd></dl></section>" for r in rows if not r["deterministic_pass"])
    leader="".join(f"<h3>{_esc(k.replace('_',' ').title())}</h3><table><tr><th>Scenario</th><th>Domain</th><th>Score</th><th>Runtime</th><th>Confidence</th></tr>{''.join(f'<tr><td>{_esc(x["scenario_id"])}</td><td>{_esc(x["domain"])}</td><td>{_esc(x["deterministic_score"])}</td><td>{_esc(x["latency_seconds"])}</td><td>{_esc(x["confidence"])}</td></tr>' for x in v)}</table>" for k,v in leaders.items())
    improvements="".join(f"<li>{_esc(x)}</li>" for x in m['top_10_improvements'])
    return f"<!doctype html><html><head><meta charset=utf-8><title>Evaluation report</title><style>body{{font:15px Arial;margin:36px;color:#0b2545}}.cards{{display:grid;grid-template-columns:repeat(3,1fr);gap:12px}}.card{{padding:15px;background:#eef4fb;border-left:5px solid #2563eb}}.card span{{display:block;font-size:20px;margin-top:6px}}img{{max-width:900px}}table{{border-collapse:collapse;width:100%;margin-bottom:20px}}th,td{{padding:7px;border-bottom:1px solid #ccd6e0;text-align:left}}pre{{white-space:pre-wrap;max-height:420px;overflow:auto;background:#f8fafc;padding:10px}}.failure{{page-break-before:always}}dt{{font-weight:bold;margin-top:8px}}</style></head><body><h1>LegacyDB Copilot Complete Local Evaluation</h1><p>Run {m['run_id']} - 125 scenarios - local MySQL</p><div class=cards>{cards}</div><h2>Executive summary</h2><p>Readiness: {m['readiness_score']}/100. Production ready: {'Yes' if m['production_ready'] else 'No'}. Operational execution completed for every case, but answer quality remains below the release threshold. The strongest result is fixture/API reliability; the largest weaknesses are exact entity resolution, evidence collection, and disabled AI reasoning.</p><h3>Biggest strengths</h3><ul>{''.join(f'<li>{_esc(x)}</li>' for x in m['biggest_strengths'])}</ul><h3>Biggest weaknesses</h3><ul>{''.join(f'<li>{_esc(x)}</li>' for x in m['biggest_weaknesses'])}</ul><h3>Top 10 improvements</h3><ol>{improvements}</ol><h2>Visualizations</h2>{imgs}<h2>Leaderboard</h2>{leader}<h2>Failed-scenario validation</h2>{failures}</body></html>"


def _docx(path,m,rows,leaders,charts):
    doc=Document(); sec=doc.sections[0]; sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
    styles=doc.styles; styles['Normal'].font.name='Calibri'; styles['Normal'].font.size=Pt(11)
    for name,size,color in [('Title',24,'0B2545'),('Heading 1',16,'2E74B5'),('Heading 2',13,'2E74B5')]:
        styles[name].font.name='Calibri'; styles[name].font.size=Pt(size); styles[name].font.color.rgb=RGBColor.from_string(color)
    doc.add_heading('LegacyDB Copilot Complete Local Evaluation',0); doc.add_paragraph(f"Run {m['run_id']} | Local MySQL | 125 scenarios")
    doc.add_heading('Executive summary',1); doc.add_paragraph(f"Readiness score: {m['readiness_score']}/100. Production ready: {'Yes' if m['production_ready'] else 'No'}. All scenarios executed, but overall accuracy was {m['overall_accuracy']:.1%} and human review was required for {m['human_review_percentage']:.1%}.")
    doc.add_heading('Biggest strengths',2)
    for value in m['biggest_strengths']: doc.add_paragraph(value,style='List Bullet')
    doc.add_heading('Biggest weaknesses',2)
    for value in m['biggest_weaknesses']: doc.add_paragraph(value,style='List Bullet')
    doc.add_heading('Top 10 improvements',2)
    for value in m['top_10_improvements']: doc.add_paragraph(value,style='List Number')
    table=doc.add_table(rows=1,cols=2); table.style='Light Shading Accent 1'; table.rows[0].cells[0].text='Metric'; table.rows[0].cells[1].text='Value'
    for k in ('overall_accuracy','precision','recall','f1_score','average_confidence','average_execution_time_seconds','average_token_usage','estimated_cost_usd','average_deterministic_score','average_ai_judge_score','pass_percentage','fail_percentage','human_review_percentage'):
        c=table.add_row().cells; c[0].text=k.replace('_',' ').title(); c[1].text=str(round(m[k],4))
    doc.add_heading('Visualizations',1)
    for name in ('overall-accuracy.png','domain-comparison.png','difficulty-comparison.png','runtime-distribution.png','confidence-distribution.png','token-usage.png','cost.png','confusion-matrix.png','pass-fail-pie.png','trend.png'):
        doc.add_picture(str(charts/name),width=Inches(6.2)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('Failed-scenario validation',1)
    for r in rows:
        if r['deterministic_pass']: continue
        doc.add_heading(f"{r['scenario_id']} - {r['severity']}",2)
        for label,key in (("Question","question"),("Expected answer","expected_answer"),("Actual answer","application_answer"),("Root cause","root_cause"),("Evidence collected","evidence_collected"),("Suggested fix","suggested_fix")):
            p=doc.add_paragraph(); p.add_run(label+': ').bold=True; p.add_run(str(r.get(key,''))[:6000])
    doc.save(path)


def _pdf(path,m,rows,leaders,charts):
    styles=getSampleStyleSheet(); story=[Paragraph('LegacyDB Copilot Complete Local Evaluation',styles['Title']),Spacer(1,12),Paragraph(f"Run {m['run_id']} | Local MySQL | 125 scenarios",styles['Normal']),Spacer(1,12)]
    story.append(Paragraph('Executive summary',styles['Heading1'])); story.append(Paragraph(f"Readiness {m['readiness_score']}/100. Production ready: {'Yes' if m['production_ready'] else 'No'}. Accuracy {m['overall_accuracy']:.1%}; human review {m['human_review_percentage']:.1%}.",styles['BodyText']))
    story.append(Paragraph('Top 10 improvements',styles['Heading2']))
    for i,value in enumerate(m['top_10_improvements'],1): story.append(Paragraph(f"{i}. {_esc(value)}",styles['BodyText']))
    metric_rows=[["Metric","Value"]]+[[k.replace('_',' ').title(),str(round(m[k],4))] for k in ('overall_accuracy','precision','recall','f1_score','average_confidence','average_execution_time_seconds','average_token_usage','estimated_cost_usd','average_deterministic_score','average_ai_judge_score','pass_percentage','fail_percentage','human_review_percentage')]
    t=Table(metric_rows,colWidths=[3.5*inch,2*inch]); t.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.HexColor('#E8EEF5')),('GRID',(0,0),(-1,-1),.3,colors.grey),('FONTSIZE',(0,0),(-1,-1),8)])); story += [Spacer(1,12),t,PageBreak()]
    for name in ('overall-accuracy.png','domain-comparison.png','difficulty-comparison.png','runtime-distribution.png','confidence-distribution.png','token-usage.png','cost.png','confusion-matrix.png','pass-fail-pie.png','trend.png'):
        story += [RLImage(str(charts/name),width=6.5*inch,height=3.5*inch),Spacer(1,8)]
    story += [PageBreak(),Paragraph('Failed-scenario validation',styles['Heading1'])]
    for r in rows:
        if r['deterministic_pass']: continue
        story += [Paragraph(f"{r['scenario_id']} - {r['severity']}",styles['Heading2'])]
        for label,key in (("Question","question"),("Expected answer","expected_answer"),("Actual answer","application_answer"),("Root cause","root_cause"),("Evidence collected","evidence_collected"),("Suggested fix","suggested_fix")):
            story += [Paragraph(f"<b>{label}:</b> {_esc(str(r.get(key,''))[:2500])}",styles['BodyText']),Spacer(1,4)]
    SimpleDocTemplate(str(path),pagesize=letter,rightMargin=.7*inch,leftMargin=.7*inch,topMargin=.7*inch,bottomMargin=.7*inch).build(story)
